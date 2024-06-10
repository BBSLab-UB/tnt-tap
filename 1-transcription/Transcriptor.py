############################################
########  TnT-TAP AUDIO TRANSCRIBER  #######
########       BBSLab Jun 2024       #######
############################################

import whisper
from datetime import datetime
import os
import requests
import io
import pandas as pd
from docx import Document
from docx.shared import Pt

start_time = datetime.now()

# MODEL LOADING
print("Loading whisper model...")
model_name = "medium"               #also: "large",...
model = whisper.load_model(model_name)
print("Whisper {} model loading completed.".format(model_name))

model_time = datetime.now()
print('Model loading lasted {}'.format(model_time - start_time))

input_ok = False

# INPUT PATH
while input_ok == False:
    input_path = os.path.normpath(input(r"Please, enter the audio or folder of audios to transcribe: ").replace("'","").replace(" ",""))
    if os.path.isdir(input_path):
        list_of_files = [file for file in os.listdir(input_path) if (os.path.isfile(os.path.join(input_path, file)) and (file[-4:] == '.WMA' or file[-4:] == '.wav'))]
        input_dir = input_path
        input_ok = True
    elif os.path.isfile(input_path):
        list_of_files = [os.path.basename(input_path),]
        input_dir = os.path.dirname(input_path)
        input_ok = True
    else:
        print("The string entered is not a valid file nor a folder!")

if len(list_of_files) == 0:
    raise Exception('No .WMA or .wav audio samples were found.')

## LANGUAGE SETTING

def check_valid_langs(language, sub):
    '''This function checks if the inputed language is supported.'''
    input_ok2 = False
    sub = sub.split('_stretch')[0]
    while input_ok2 == False:
        if language.capitalize() in langs_whisper:
            input_ok2 = True
        else:
            language = input("Audio language '{}' for subject {} is not supported on Whisper. Please enter, in English, the transcription language: ".format(language.capitalize(), sub))
    return language

# Supported languages whisper
langs_whisper = [
         'Afrikaans', 'Albanian', 'Amharic', 'Arabic', 'Armenian', 'Assamese', 
         'Azerbaijani', 'Bashkir', 'Basque', 'Belarusian', 'Bengali', 
         'Bosnian', 'Breton', 'Bulgarian', 'Burmese', 'Cantonese', 'Castilian',
         'Catalan', 'Chinese', 'Croatian', 'Czech', 'Danish', 'Dutch', 
         'English', 'Estonian', 'Faroese', 'Finnish', 'Flemish', 'French', 
         'Galician', 'Georgian', 'German', 'Greek', 'Gujarati', 'Haitian', 
         'Haitian Creole', 'Hausa', 'Hawaiian', 'Hebrew', 'Hindi', 'Hungarian',
         'Icelandic', 'Indonesian', 'Italian', 'Japanese', 'Javanese', 
         'Kannada', 'Kazakh', 'Khmer', 'Korean', 'Lao', 'Latin', 'Latvian', 
         'Letzeburgesch', 'Lingala', 'Lithuanian', 'Luxembourgish', 
         'Macedonian', 'Malagasy', 'Malay', 'Malayalam', 'Maltese', 'Mandarin',
         'Maori', 'Marathi', 'Moldavian', 'Moldovan', 'Mongolian', 'Myanmar', 
         'Nepali', 'Norwegian', 'Nynorsk', 'Occitan', 'Panjabi', 'Pashto', 
         'Persian', 'Polish', 'Portuguese', 'Punjabi', 'Pushto', 'Romanian', 
         'Russian', 'Sanskrit', 'Serbian', 'Shona', 'Sindhi', 'Sinhala', 
         'Sinhalese', 'Slovak', 'Slovenian', 'Somali', 'Spanish', 'Sundanese', 
         'Swahili', 'Swedish', 'Tagalog', 'Tajik', 'Tamil', 'Tatar', 'Telugu', 
         'Thai', 'Tibetan', 'Turkish', 'Turkmen', 'Ukrainian', 'Urdu', 'Uzbek',
         'Valencian', 'Vietnamese', 'Welsh', 'Yiddish', 'Yoruba']

# GET PRIVATE API KEY
with open(os.path.join(os.path.dirname(os.path.dirname(__file__)), 'api_key.txt'), 'r') as f:
    key = f.read().replace('\n', '')

# DATA DEFINITION FOR REDCAP API
data = {
    'token': key,
    'content': 'report',
    'format': 'csv',
    'report_id': '23',
    'rawOrLabel': 'label',
    'returnFormat': 'json'
}

#REDCAP API ACCESS
r = requests.post(url = 'http://www.ub.edu/neuropsicologia/redcap/api/', data = data)  
if 'error' in r.text:
    raise ValueError(r.text)
    
df = pd.read_csv(io.StringIO(r.text), header=0)
df = df.rename({'record_id_np_w2': 'SUBJECT', 'tap_language_w2': 'LANGUAGE'}, axis=1)
df.SUBJECT = df.SUBJECT.astype(str)
df = df.set_index('SUBJECT')

len_list = len(list_of_files)
len_list0 = len(list_of_files)

# Dataframe to dictionary
list_of_subs = [str(os.path.splitext(filename)[0]).split('_stretch')[0] for filename in list_of_files]
list_of_subs_in_df = [sub for sub in list_of_subs if sub in df.index.values]
dict_subs_langs = list(df.loc[list_of_subs_in_df].to_dict().values())[0]
for sub in set(list_of_subs).difference(set(list_of_subs_in_df)):
    dict_subs_langs[sub] = ''

# Input missing values
keys = list(dict_subs_langs.keys())
values = list(dict_subs_langs.values())
for idx, lang in enumerate(values):
    # print(idx, lang)
    sub = keys[idx]
    if pd.isna(lang) or lang == '':
        lang = input("Audio language for subject {} was not found in REDCap. Please enter, in English, the transcription language: ".format(sub))
        lang = check_valid_langs(lang, sub)
        dict_subs_langs[sub] = lang.capitalize()
    lang = check_valid_langs(lang, sub)        
    dict_subs_langs[sub] = lang.capitalize()

# Language check
input_ok4 = False
while input_ok4 == False:
    df_subs_langs = pd.DataFrame.from_dict(dict_subs_langs, orient='index', columns=['LANGUAGE',])
    df_subs_langs.index.name = 'SUBJECT'
    print("Subjects will be transcribed in the following languages: \n")
    print(df_subs_langs)
    
    input_ok3 = False
    while input_ok3 == False:
        quest_lang = input("\nAre all languages correct? (Y/N) ").upper()
        if quest_lang == 'Y':
            input_ok3 = True
            input_ok4 = True
        elif quest_lang == 'N':
            if len(dict_subs_langs) > 1:
                sub_mod = input("Please enter the ID of the subject whose language you want to change: ")
                while sub_mod not in keys:
                    sub_mod = input("Sub ID {} was not found in the list of subjects to transcribe. Please enter a valid ID from the list {}: ".format(sub_mod, keys))
            else:
                sub_mod = list(dict_subs_langs)[0]
            new_lang = input("Please enter, in English, the transcription language for subject {}: ".format(sub_mod))
            new_lang = check_valid_langs(new_lang, sub_mod)
            dict_subs_langs[sub_mod] = new_lang.capitalize()
            input_ok3 = True
        else:
            print("Please, enter a valid response.")

# TRANSCRIPTION
for filename in list_of_files:
    fpath = os.path.join(input_dir, filename)
    subject = str(os.path.splitext(filename)[0])
    sub_id = subject.split('_stretch')[0]
    language = dict_subs_langs[sub_id]
    print('---- TRANSCRIPTION OF SUBJECT {} ----'.format(sub_id))
            
    # TRANSCRIPTION
    transcribe_time0 = datetime.now()
    print("Starting transcription for subject {}. {} out of {} subjects left.".format(sub_id, len_list, len_list0))
    result = model.transcribe(fpath, language = language, fp16 = False, verbose=False)
    transcribe_time = datetime.now()
    print('Transcription of document {} lasted {}'.format(os.path.basename(fpath), transcribe_time - transcribe_time0))
    
    print("Generating document...")
    transcripts_dir = os.path.join(os.path.dirname(fpath), 'Transcripts')
    os.makedirs(transcripts_dir, exist_ok=True)

    # DOCUMENT CREATION
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = 1

    document.add_paragraph(subject)
    document.add_paragraph("")
    text = result['text']

    if text[0] == " ":
        text=text[1:]
    
    document.add_paragraph(text)
    output_path = os.path.join(os.path.dirname(fpath), 'Transcripts', subject + '_transcript.docx')
    document.save(output_path)
    doc_time = datetime.now()
    print('{} document creation lasted {}'.format(output_path, doc_time - transcribe_time))
    
    len_list = len_list-1